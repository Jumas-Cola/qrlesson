from django.shortcuts import render
from django.shortcuts import get_object_or_404
from django.http import HttpResponseRedirect, HttpResponse
from .models import TeachingMethod, LessonPart
from django.views import generic
from docx import Document
import os
import json
from datetime import date
from django.urls import reverse_lazy


def index(request):

    return render(
        request,
        'index.html',
        context={
            'section_list': LessonPart.objects.all(),
        },
    )


def about(request):

    return render(
        request,
        'about.html',
        context={
            'section_list': LessonPart.objects.all(),
        },
    )


def authors(request):

    return render(
        request,
        'authors.html',
        context={
            'section_list': LessonPart.objects.all(),
        },
    )


def qr_table(request):

    return render(
        request,
        'qr_table.html',
        context={
            'section_list': LessonPart.objects.all(),
        },
    )


def qr_generator(request):

    return render(
        request,
        'qr_generator.html',
        context={
            'section_list': LessonPart.objects.all(),
        },
    )


def chosen_methods(request):
    chosen_methods = []
    if 'methods' not in request.session:
        request.session['methods'] = []
    for id in request.session['methods']:
        chosen_methods.append(TeachingMethod.objects.filter(id=id)[0])

    return render(
        request,
        'qrmethods/chosen_methods.html',
        context={
            'chosen_methods': chosen_methods,
            'section_list': LessonPart.objects.all(),
        },
    )


def upd_session(request, pk):
    if 'methods' in request.session:
        if (str(pk) not in request.session['methods']):
            request.session['methods'] += [pk]
    else:
        request.session['methods'] = [pk]
    return HttpResponseRedirect(request.GET['next'])


def rmv_from_session(request):
    ids_json = request.POST.get('ids')
    method = request.POST.get('method')
    ids = json.loads(ids_json)
    request.session['methods'] = [x for x in ids if x != method]
    return HttpResponseRedirect(reverse_lazy('chosen_methods'))


class MethodListView(generic.ListView):
    model = TeachingMethod
    template_name = 'qrmethods/method_list.html'
    paginate_by = 9

    def get_queryset(self):
        qs = super().get_queryset()
        lp = self.kwargs['section']
        get_object_or_404(LessonPart, url=lp)
        return qs.filter(lesson_part__url=lp)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(section_list=LessonPart.objects.all(), **kwargs)
        context['section'] = LessonPart.objects.filter(
            url=self.kwargs['section'])[0]
        return context


class MethodDetailView(generic.DetailView):
    model = TeachingMethod
    template_name = 'qrmethods/method_detail.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(section_list=LessonPart.objects.all(), **kwargs)
        if 'methods' not in self.request.session:
            self.request.session['methods'] = []
        context['methods'] = self.request.session['methods']
        if 'next' in self.request.GET:
            context['next'] = self.request.GET['next']
            self.request.session['next'] = self.request.GET['next']
        elif 'next' in self.request.session:
            context['next'] = self.request.session['next']
        return context

    # def get_object(self, queryset=None):
    #     obj = super().get_object()
    #     if not obj.author == self.request.user:
    #         raise Http404
    #     return obj


def download(request):
    ids_json = request.POST.get('ids')
    ids = json.loads(ids_json)
    document = Document()
    docx_title = "lesson_plan_{}.docx".format(date.today())
    # ---- Cover Letter ----
    document.add_heading('План урока', 0)

    for id in ids:
        method = TeachingMethod.objects.get(pk=id)
        document.add_heading(method.title, level=2)
        document.add_paragraph(method.description)
    document.add_page_break()
    document.save(docx_title)

    file_path = os.path.realpath(docx_title)
    data = open(file_path, "rb").read()
    response = HttpResponse(
        data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = os.path.getsize(file_path)
    os.remove(file_path)

    return response
